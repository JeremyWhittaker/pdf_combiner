"""Pytest configuration and fixtures."""

import tempfile
from pathlib import Path
from typing import Generator

import pytest
from PyPDF2 import PdfWriter

from pdf_combiner.config import Config
from pdf_combiner.models import DocumentInfo, DocumentType


@pytest.fixture
def temp_dir() -> Generator[Path, None, None]:
    """Create a temporary directory for tests."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def sample_pdf(temp_dir: Path) -> Path:
    """Create a sample PDF file."""
    pdf_path = temp_dir / "sample.pdf"
    
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    
    with open(pdf_path, "wb") as f:
        writer.write(f)
    
    return pdf_path


@pytest.fixture
def sample_pdf_with_text(temp_dir: Path) -> Path:
    """Create a sample PDF file with text."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    pdf_path = temp_dir / "sample_with_text.pdf"
    
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(100, 750, "This is a test PDF with text content.")
    c.drawString(100, 700, "It should not require OCR processing.")
    c.showPage()
    c.save()
    
    return pdf_path


@pytest.fixture
def sample_docx(temp_dir: Path) -> Path:
    """Create a sample DOCX file."""
    from docx import Document
    
    docx_path = temp_dir / "sample.docx"
    
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test DOCX file.")
    doc.add_paragraph("It contains some sample text.")
    doc.save(str(docx_path))
    
    return docx_path


@pytest.fixture
def config() -> Config:
    """Create a test configuration."""
    return Config(
        output={
            "default_name": "test_combined.pdf",
            "add_metadata": True,
            "compression": True,
            "overwrite": True,
        },
        ocr={
            "enabled": False,  # Disable OCR for most tests
            "language": "eng",
            "dpi": 300,
            "skip_text_pages": True,
            "timeout": 60,
        },
        processing={
            "max_workers": 2,
            "batch_size": 5,
            "fail_fast": False,
        },
        logging={
            "level": "DEBUG",
            "format": "%(asctime)s [%(levelname)s] %(message)s",
        }
    )


@pytest.fixture
def sample_documents(sample_pdf: Path, sample_docx: Path) -> list[DocumentInfo]:
    """Create a list of sample documents."""
    docs = []
    
    for path in [sample_pdf, sample_docx]:
        stat = path.stat()
        doc_type = DocumentType.PDF if path.suffix == ".pdf" else DocumentType.DOCX
        
        doc = DocumentInfo(
            path=path,
            name=path.name,
            type=doc_type,
            size_bytes=stat.st_size,
            created_at=stat.st_ctime,
            modified_at=stat.st_mtime,
        )
        docs.append(doc)
    
    return docs


@pytest.fixture
def mock_ocrmypdf(monkeypatch):
    """Mock ocrmypdf command."""
    def mock_run(cmd, *args, **kwargs):
        # Simulate ocrmypdf by copying input to output
        if "ocrmypdf" in cmd[0]:
            input_file = Path(cmd[-2])
            output_file = Path(cmd[-1])
            if input_file.exists():
                import shutil
                shutil.copy(input_file, output_file)
        
        from subprocess import CompletedProcess
        return CompletedProcess(cmd, 0, "", "")
    
    import subprocess
    monkeypatch.setattr(subprocess, "run", mock_run)


@pytest.fixture
def mock_libreoffice(monkeypatch):
    """Mock LibreOffice command."""
    def mock_run(cmd, *args, **kwargs):
        # Simulate LibreOffice conversion
        if "libreoffice" in cmd[0]:
            for i, arg in enumerate(cmd):
                if arg.endswith((".doc", ".docx")):
                    input_file = Path(arg)
                    output_dir = Path(cmd[cmd.index("--outdir") + 1])
                    output_file = output_dir / f"{input_file.stem}.pdf"
                    
                    # Create a simple PDF
                    from PyPDF2 import PdfWriter
                    writer = PdfWriter()
                    writer.add_blank_page(width=200, height=200)
                    with open(output_file, "wb") as f:
                        writer.write(f)
                    break
        
        from subprocess import CompletedProcess
        return CompletedProcess(cmd, 0, "", "")
    
    import subprocess
    monkeypatch.setattr(subprocess, "run", mock_run)